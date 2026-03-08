#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
跨学科项目化学习设计方案 Word 文档生成器

依赖：pip install python-docx
"""

import argparse
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("错误：缺少 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


class PBLDocGenerator:
    """PBL设计方案文档生成器"""
    
    def __init__(self, template_type='detailed'):
        self.doc = Document()
        self.template_type = template_type
        self._setup_styles()
        
    def _setup_styles(self):
        """设置文档样式"""
        # 设置中文字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)
        
        # 标题样式
        for i in range(1, 5):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = '黑体'
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            if i == 1:
                heading_style.font.size = Pt(18)
            elif i == 2:
                heading_style.font.size = Pt(16)
            elif i == 3:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)
    
    def add_cover_page(self, project_info):
        """添加封面页"""
        # 标题
        title = self.doc.add_heading(f'{project_info.get("project_name", "[项目名称]")}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # 副标题
        subtitle = self.doc.add_paragraph('跨学科项目化学习设计方案')
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        
        # 空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 项目信息表格
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info_items = [
            ('适用年级', project_info.get('grade', '[年级]')),
            ('涉及学科', project_info.get('subjects', '[学科]')),
            ('项目时长', project_info.get('duration', '[周数/课时]')),
            ('设计者', project_info.get('designer', '[姓名]')),
            ('设计日期', project_info.get('date', datetime.now().strftime('%Y年%m月%d日')))
        ]
        
        for i, (label, value) in enumerate(info_items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            
        # 分页
        self.doc.add_page_break()
    
    def add_section_project_basis(self, content):
        """添加"项目依据"部分"""
        self.doc.add_heading('一、项目依据', 1)
        
        # 1.1 真实世界情境
        self.doc.add_heading('1.1 真实世界情境', 2)
        para = self.doc.add_paragraph(content.get('real_world_context', 
            '[描述项目如何源于真实生活问题、社会需求或学生关注的议题...]'))
        
        # 1.2 课程标准对接
        self.doc.add_heading('1.2 课程标准对接', 2)
        
        curriculum_standards = content.get('curriculum_standards', {})
        if curriculum_standards:
            for subject, standards in curriculum_standards.items():
                self.doc.add_paragraph(f'{subject}课程标准：', style='Heading 3')
                for standard in standards:
                    self.doc.add_paragraph(standard, style='List Bullet')
        else:
            self.doc.add_paragraph('[各学科课程标准要求...]')
        
        # 1.3 学生学情分析
        self.doc.add_heading('1.3 学生学情分析', 2)
        para = self.doc.add_paragraph(content.get('student_analysis', 
            '[分析目标年级学生的认知特点、已有基础、兴趣特点和发展需求...]'))
    
    def add_section_project_overview(self, content):
        """添加"项目概述"部分"""
        self.doc.add_heading('二、项目概述', 1)
        
        overview_items = [
            ('项目名称', content.get('project_name', '[项目名称]')),
            ('适用年级', content.get('grade', '[年级]')),
            ('涉及学科', content.get('subjects', '[学科及占比]')),
            ('项目时长', content.get('duration', '[周数或课时]')),
            ('项目简介', content.get('description', '[3-5句话概括项目内容和意义]'))
        ]
        
        for label, value in overview_items:
            para = self.doc.add_paragraph()
            para.add_run(f'• {label}：').bold = True
            para.add_run(value)
    
    def add_section_learning_objectives(self, content):
        """添加"学习目标与核心素养"部分"""
        self.doc.add_heading('三、学习目标与核心素养', 1)
        
        # 3.1 学科知识目标
        self.doc.add_heading('3.1 学科知识目标', 2)
        knowledge_goals = content.get('knowledge_goals', {})
        if knowledge_goals:
            for subject, goals in knowledge_goals.items():
                self.doc.add_paragraph(f'{subject}：', style='Heading 3')
                for goal in goals:
                    self.doc.add_paragraph(goal, style='List Bullet')
        else:
            self.doc.add_paragraph('[各学科具体知识目标...]')
        
        # 3.2 跨学科能力目标
        self.doc.add_heading('3.2 跨学科能力目标', 2)
        ability_goals = content.get('ability_goals', [
            '[问题解决能力目标]',
            '[信息处理能力目标]',
            '[实践创新能力目标]'
        ])
        for goal in ability_goals:
            self.doc.add_paragraph(goal, style='List Bullet')
        
        # 3.3 核心素养发展目标
        self.doc.add_heading('3.3 核心素养发展目标', 2)
        competency_goals = content.get('competency_goals', [
            '批判性思维：[具体表现]',
            '创造力：[具体表现]',
            '合作能力：[具体表现]',
            '沟通表达：[具体表现]'
        ])
        for goal in competency_goals:
            self.doc.add_paragraph(goal, style='List Bullet')
    
    def add_section_driving_questions(self, content):
        """添加"驱动性问题"部分"""
        self.doc.add_heading('四、驱动性问题', 1)
        
        questions = content.get('driving_questions', {})
        
        # 核心问题
        core_question = questions.get('core', '[核心驱动性问题]')
        para = self.doc.add_paragraph()
        para.add_run('核心问题：').bold = True
        para.add_run(core_question)
        
        # 子问题
        sub_questions = questions.get('sub_questions', [
            '[子问题1]',
            '[子问题2]',
            '[子问题3]'
        ])
        
        self.doc.add_paragraph()
        para = self.doc.add_paragraph()
        para.add_run('子问题：').bold = True
        
        for i, question in enumerate(sub_questions, 1):
            self.doc.add_paragraph(f'{i}. {question}', style='List Number')
    
    def add_section_learning_activities(self, content):
        """添加"学习活动设计"部分"""
        self.doc.add_heading('五、学习活动设计', 1)
        
        activities = content.get('activities', [])
        
        if not activities:
            # 默认模板
            activities = [
                {
                    'name': '阶段一：[阶段名称]',
                    'duration': 'X课时',
                    'objectives': '[学习目标]',
                    'main_activities': '[主要活动内容]',
                    'subject_integration': {'学科1': '[融合点]', '学科2': '[融合点]'},
                    'outcomes': '[学习成果]',
                    'teacher_guidance': '[教师指导要点]'
                }
            ]
        
        for i, activity in enumerate(activities, 1):
            self.doc.add_heading(f'阶段{["一", "二", "三", "四", "五"][i-1]}：{activity.get("name", "")}（{activity.get("duration", "X课时")}）', 2)
            
            # 学习目标
            para = self.doc.add_paragraph()
            para.add_run('学习目标：').bold = True
            para.add_run(activity.get('objectives', '[目标]'))
            
            # 主要活动
            para = self.doc.add_paragraph()
            para.add_run('主要活动：').bold = True
            para.add_run(activity.get('main_activities', '[活动内容]'))
            
            # 学科融合点
            para = self.doc.add_paragraph()
            para.add_run('学科融合点：').bold = True
            subject_integration = activity.get('subject_integration', {})
            for subject, integration in subject_integration.items():
                self.doc.add_paragraph(f'• {subject}：{integration}')
            
            # 学习成果
            para = self.doc.add_paragraph()
            para.add_run('学习成果：').bold = True
            para.add_run(activity.get('outcomes', '[成果]'))
            
            # 教师指导
            para = self.doc.add_paragraph()
            para.add_run('教师指导：').bold = True
            para.add_run(activity.get('teacher_guidance', '[指导要点]'))
            
            self.doc.add_paragraph()  # 空行
    
    def add_section_assessment(self, content):
        """添加"评估方案"部分"""
        self.doc.add_heading('六、评估方案', 1)
        
        # 6.1 形成性评估
        self.doc.add_heading('6.1 形成性评估', 2)
        
        formative = content.get('formative_assessment', [
            {'method': '学习日志', 'frequency': '每周', 'focus': '反思与思考'},
            {'method': '小组讨论观察', 'frequency': '每次课', 'focus': '参与度与贡献'},
            {'method': '阶段性成果检查', 'frequency': '每阶段', 'focus': '任务完成质量'}
        ])
        
        table = self.doc.add_table(rows=len(formative)+1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        table.rows[0].cells[0].text = '评估方式'
        table.rows[0].cells[1].text = '频率'
        table.rows[0].cells[2].text = '评估重点'
        
        for i, item in enumerate(formative, 1):
            table.rows[i].cells[0].text = item.get('method', '')
            table.rows[i].cells[1].text = item.get('frequency', '')
            table.rows[i].cells[2].text = item.get('focus', '')
        
        self.doc.add_paragraph()
        
        # 6.2 总结性评估
        self.doc.add_heading('6.2 总结性评估', 2)
        
        summative = content.get('summative_assessment', [
            {'content': '学科知识掌握', 'weight': '40%', 'criteria': '[具体标准]'},
            {'content': '问题解决能力', 'weight': '30%', 'criteria': '[具体标准]'},
            {'content': '合作与沟通', 'weight': '20%', 'criteria': '[具体标准]'},
            {'content': '创新与反思', 'weight': '10%', 'criteria': '[具体标准]'}
        ])
        
        table = self.doc.add_table(rows=len(summative)+1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = '评估内容'
        table.rows[0].cells[1].text = '权重'
        table.rows[0].cells[2].text = '评分标准'
        
        for i, item in enumerate(summative, 1):
            table.rows[i].cells[0].text = item.get('content', '')
            table.rows[i].cells[1].text = item.get('weight', '')
            table.rows[i].cells[2].text = item.get('criteria', '')
        
        self.doc.add_paragraph()
        
        # 6.3 评分量规
        self.doc.add_heading('6.3 评分量规（Rubric）', 2)
        self.doc.add_paragraph('[详见附录：4级评分标准表]')
    
    def add_section_resources(self, content):
        """添加"所需资源与材料"部分"""
        self.doc.add_heading('七、所需资源与材料', 1)
        
        resources = content.get('resources', {})
        
        resource_types = [
            ('materials', '7.1 材料资源'),
            ('technology', '7.2 技术资源'),
            ('learning', '7.3 学习资源'),
            ('human_venue', '7.4 人力与场地资源')
        ]
        
        for key, heading in resource_types:
            self.doc.add_heading(heading, 2)
            items = resources.get(key, ['[资源项目：用途、数量]'])
            for item in items:
                self.doc.add_paragraph(item, style='List Bullet')
    
    def add_section_timeline(self, content):
        """添加"项目时间线"部分"""
        self.doc.add_heading('八、项目时间线', 1)
        
        timeline = content.get('timeline', [
            {'week': '第1周', 'stage': '启动', 'tasks': '[主要任务]', 'outcomes': '[预期成果]'},
            {'week': '第2周', 'stage': '探究', 'tasks': '[主要任务]', 'outcomes': '[预期成果]'},
            {'week': '第3周', 'stage': '制作', 'tasks': '[主要任务]', 'outcomes': '[预期成果]'},
            {'week': '第4周', 'stage': '展示', 'tasks': '[主要任务]', 'outcomes': '[预期成果]'}
        ])
        
        table = self.doc.add_table(rows=len(timeline)+1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = '周次'
        table.rows[0].cells[1].text = '阶段'
        table.rows[0].cells[2].text = '主要任务'
        table.rows[0].cells[3].text = '预期成果'
        
        for i, item in enumerate(timeline, 1):
            table.rows[i].cells[0].text = item.get('week', '')
            table.rows[i].cells[1].text = item.get('stage', '')
            table.rows[i].cells[2].text = item.get('tasks', '')
            table.rows[i].cells[3].text = item.get('outcomes', '')
    
    def add_section_differentiation(self, content):
        """添加"差异化教学策略"部分"""
        self.doc.add_heading('九、差异化教学策略', 1)
        
        # 9.1 学习支持
        self.doc.add_heading('9.1 学习支持', 2)
        support = content.get('support', ['[为学习困难学生提供的支持和脚手架]'])
        for item in support:
            self.doc.add_paragraph(item, style='List Bullet')
        
        # 9.2 拓展挑战
        self.doc.add_heading('9.2 拓展挑战', 2)
        challenge = content.get('challenge', ['[为优秀学生设计的进阶任务]'])
        for item in challenge:
            self.doc.add_paragraph(item, style='List Bullet')
        
        # 9.3 灵活分组
        self.doc.add_heading('9.3 灵活分组策略', 2)
        grouping = content.get('grouping', '[分组原则和调整机制]')
        self.doc.add_paragraph(grouping)
    
    def add_section_implementation(self, content):
        """添加"实施建议"部分"""
        self.doc.add_heading('十、实施建议', 1)
        
        # 10.1 教师准备
        self.doc.add_heading('10.1 教师准备', 2)
        preparation = content.get('teacher_preparation', ['[需要提前准备的内容]'])
        for item in preparation:
            self.doc.add_paragraph(item, style='List Bullet')
        
        # 10.2 常见问题与应对
        self.doc.add_heading('10.2 常见问题与应对', 2)
        problems = content.get('common_problems', ['[可能遇到的挑战及解决方案]'])
        for item in problems:
            self.doc.add_paragraph(item, style='List Bullet')
        
        # 10.3 家长沟通
        self.doc.add_heading('10.3 家长沟通', 2)
        parent_communication = content.get('parent_communication', '[如何争取家长理解和支持]')
        self.doc.add_paragraph(parent_communication)
    
    def add_appendix(self):
        """添加附录"""
        self.doc.add_page_break()
        self.doc.add_heading('附录', 1)
        
        self.doc.add_heading('附录A：学习单模板', 2)
        self.doc.add_paragraph('[具体的学习单模板]')
        
        self.doc.add_heading('附录B：评估表格', 2)
        self.doc.add_paragraph('[具体的评估表格]')
        
        self.doc.add_heading('附录C：参考资料清单', 2)
        self.doc.add_paragraph('[推荐的参考书籍、网站、视频等]')
    
    def generate(self, content, output_path):
        """生成完整文档"""
        try:
            # 封面
            self.add_cover_page(content.get('project_info', {}))
            
            # 各部分内容
            self.add_section_project_basis(content.get('project_basis', {}))
            self.add_section_project_overview(content.get('project_overview', {}))
            self.add_section_learning_objectives(content.get('learning_objectives', {}))
            self.add_section_driving_questions(content.get('driving_questions', {}))
            self.add_section_learning_activities(content.get('learning_activities', {}))
            self.add_section_assessment(content.get('assessment', {}))
            self.add_section_resources(content.get('resources', {}))
            self.add_section_timeline(content.get('timeline', {}))
            self.add_section_differentiation(content.get('differentiation', {}))
            self.add_section_implementation(content.get('implementation', {}))
            self.add_appendix()
            
            # 保存
            self.doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"生成文档时出错：{e}")
            return False


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='生成跨学科项目化学习设计方案Word文档'
    )
    parser.add_argument(
        '--output', '-o',
        default='PBL项目设计方案.docx',
        help='输出文件名（默认：PBL项目设计方案.docx）'
    )
    parser.add_argument(
        '--template',
        choices=['detailed', 'concise'],
        default='detailed',
        help='模板类型（默认：detailed）'
    )
    parser.add_argument(
        '--open',
        action='store_true',
        help='生成后自动打开文档'
    )
    
    args = parser.parse_args()
    
    print("正在生成PBL项目设计方案文档...")
    
    # 创建生成器
    generator = PBLDocGenerator(template_type=args.template)
    
    # 默认内容（AI助手会替换这些内容）
    default_content = {
        'project_info': {
            'project_name': '[项目名称]',
            'grade': '[年级]',
            'subjects': '[学科]',
            'duration': '[时长]',
            'designer': '[设计者]',
            'date': datetime.now().strftime('%Y年%m月%d日')
        }
    }
    
    # 生成文档
    output_path = Path(args.output)
    success = generator.generate(default_content, output_path)
    
    if success:
        print(f"✓ 文档生成成功：{output_path.absolute()}")
        
        if args.open:
            import platform
            import subprocess
            
            system = platform.system()
            try:
                if system == 'Darwin':  # macOS
                    subprocess.run(['open', str(output_path)])
                elif system == 'Windows':
                    subprocess.run(['start', str(output_path)], shell=True)
                elif system == 'Linux':
                    subprocess.run(['xdg-open', str(output_path)])
                print("✓ 已打开文档")
            except Exception as e:
                print(f"无法自动打开文档：{e}")
    else:
        print("✗ 文档生成失败")
        sys.exit(1)


if __name__ == '__main__':
    main()
